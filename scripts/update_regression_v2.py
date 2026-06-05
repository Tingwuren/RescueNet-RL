"""
修改 13 需求回归分析.docx：
1. 保持封面和目录完全不变
2. 将正文中"广播通信融合网络模拟与资源优化调度仿真系统（原型系统）"
   统一加书名号为《广播通信融合网络模拟与资源优化调度仿真系统（原型系统）》
3. 根据20260530会议纪要，对仿真系统相关内容进行补充
"""

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy, re

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
BASE = '/mnt/data0/root/Projects/RescueNet-RL/word/中期反馈/修改文件/'

# ── 格式构建工具 ──────────────────────────────────────────────────────────────

def _rpr(bold=False):
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(f'{{{NS}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS}}}eastAsia', '宋体')
    rFonts.set(f'{{{NS}}}cs', 'Times New Roman')
    rPr.append(rFonts)
    b = OxmlElement('w:b')
    if not bold:
        b.set(f'{{{NS}}}val', '0')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(f'{{{NS}}}val', '24')
    rPr.append(sz)
    return rPr

def _run(text, bold=False):
    r = OxmlElement('w:r')
    r.append(_rpr(bold))
    t = OxmlElement('w:t')
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r

def make_body_para(text, indent=True):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(f'{{{NS}}}before', '0')
    sp.set(f'{{{NS}}}after', '0')
    sp.set(f'{{{NS}}}line', '360')
    sp.set(f'{{{NS}}}lineRule', 'auto')
    pPr.append(sp)
    ind = OxmlElement('w:ind')
    ind.set(f'{{{NS}}}left', '0')
    ind.set(f'{{{NS}}}right', '0')
    ind.set(f'{{{NS}}}firstLine', '480' if indent else '0')
    pPr.append(ind)
    p.append(pPr)
    p.append(_run(text))
    return p

def make_h2(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    ps = OxmlElement('w:pStyle'); ps.set(f'{{{NS}}}val', '3'); pPr.append(ps)
    np = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(f'{{{NS}}}val', '1'); np.append(il)
    ni = OxmlElement('w:numId'); ni.set(f'{{{NS}}}val', '0'); np.append(ni)
    pPr.append(np)
    sp = OxmlElement('w:spacing')
    sp.set(f'{{{NS}}}before', '60'); sp.set(f'{{{NS}}}after', '60')
    sp.set(f'{{{NS}}}line', '360'); sp.set(f'{{{NS}}}lineRule', 'auto')
    pPr.append(sp)
    ind = OxmlElement('w:ind')
    ind.set(f'{{{NS}}}left', '0'); ind.set(f'{{{NS}}}right', '0')
    ind.set(f'{{{NS}}}firstLine', '0'); pPr.append(ind)
    p.append(pPr)
    p.append(_run(text, bold=True))
    return p

def insert_after(body, ref_el, new_el):
    children = list(body)
    for i, c in enumerate(children):
        if c is ref_el:
            body.insert(i + 1, new_el)
            return i + 1
    return -1

def insert_before(body, ref_el, new_el):
    children = list(body)
    for i, c in enumerate(children):
        if c is ref_el:
            body.insert(i, new_el)
            return i
    return -1

def find_para_el(body, text_contains):
    for child in list(body):
        if child.tag == f'{{{NS}}}p':
            t_els = child.findall(f'.//{{{NS}}}t')
            text = ''.join(t.text or '' for t in t_els)
            if text_contains in text:
                return child
    return None

def find_all_para_els(body, text_contains):
    results = []
    for child in list(body):
        if child.tag == f'{{{NS}}}p':
            t_els = child.findall(f'.//{{{NS}}}t')
            text = ''.join(t.text or '' for t in t_els)
            if text_contains in text:
                results.append(child)
    return results

def get_para_text(el):
    t_els = el.findall(f'.//{{{NS}}}t')
    return ''.join(t.text or '' for t in t_els)

def replace_run_text_in_para(para_el, old, new):
    """在段落的所有 run 的文本中替换 old→new（可能跨 run，直接替换拼合文本后重建）"""
    runs = para_el.findall(f'{{{NS}}}r')
    if not runs:
        return
    # 拼合全文
    full = ''.join((t.text or '') for r in runs for t in r.findall(f'{{{NS}}}t'))
    if old not in full:
        return
    new_full = full.replace(old, new)
    # 保留第一个 run 的格式，清空其余 run
    first_r = runs[0]
    for r in runs[1:]:
        para_el.remove(r)
    t_els = first_r.findall(f'{{{NS}}}t')
    for t in t_els:
        first_r.remove(t)
    t = OxmlElement('w:t')
    t.text = new_full
    first_r.append(t)

# ── 书名号替换 ─────────────────────────────────────────────────────────────────
SYSTEM_NAME = '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）'
SYSTEM_NAME_WITH_BRACKETS = f'《{SYSTEM_NAME}》'

# 需要加书名号的文档名列表
DOC_NAMES = [
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）总体设计方案',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）开发手册',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）UI设计文档',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）设备类型字典',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）验收测试大纲',
    '广播通信融合网络模拼与资源优化调度仿真系统（原型系统）中期指标自测报告',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）中期指标测试大纲',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）中期指标自测报告',
    '广播通信融合网络模拼与资源优化调度仿真系统（原型系统）当前系统实现材料',
    '广播通信融合网络模拟与资源优化调度仿真系统（原型系统）当前系统实现材料',
]

def add_brackets_to_doc_names(para_el):
    """给段落中出现的文档名加书名号（如果尚未有书名号）"""
    runs = para_el.findall(f'{{{NS}}}r')
    if not runs:
        return
    full = ''.join((t.text or '') for r in runs for t in r.findall(f'{{{NS}}}t'))
    changed = False
    for doc_name in DOC_NAMES:
        # 如果已经有书名号则跳过
        if f'《{doc_name}》' in full:
            continue
        if doc_name in full:
            full = full.replace(doc_name, f'《{doc_name}》')
            changed = True
    if not changed:
        return
    # 重建 runs
    first_r = runs[0]
    for r in runs[1:]:
        para_el.remove(r)
    for t in first_r.findall(f'{{{NS}}}t'):
        first_r.remove(t)
    t = OxmlElement('w:t')
    t.text = full
    first_r.append(t)


# ── 主程序 ────────────────────────────────────────────────────────────────────

doc = docx.Document(BASE + '13 需求回归分析.docx')
body = doc.element.body
paras = doc.paragraphs

# 确定封面/目录结束位置（第一个 Heading 1 之前）
def get_pStyle(el):
    s = el.find(f'.//{{{NS}}}pStyle')
    return s.get(f'{{{NS}}}val') if s is not None else None

first_h1_el = None
for child in list(body):
    if child.tag == f'{{{NS}}}p' and get_pStyle(child) == '2':
        first_h1_el = child
        break

cover_end_idx = list(body).index(first_h1_el)
print(f'封面/目录结束于XML子元素索引 {cover_end_idx}，第一个H1: {get_para_text(first_h1_el)[:30]}')

# ── STEP 1: 书名号替换（仅正文部分，索引 >= cover_end_idx）────────────────────
body_children = list(body)
for i, child in enumerate(body_children):
    if i < cover_end_idx:
        continue
    if child.tag == f'{{{NS}}}p':
        add_brackets_to_doc_names(child)
    elif child.tag == f'{{{NS}}}tbl':
        for tc in child.findall(f'.//{{{NS}}}tc'):
            for p_el in tc.findall(f'{{{NS}}}p'):
                add_brackets_to_doc_names(p_el)

print('书名号替换完成')

# ── STEP 2: 补充"1 背景"内容 ─────────────────────────────────────────────────
# 找到第47段（"中期反馈后..."）替换为包含会议背景的更完整版本
# 先重新获取段落列表
doc.save(BASE + '13 需求回归分析_new.docx')
doc = docx.Document(BASE + '13 需求回归分析_new.docx')
body = doc.element.body

# 找"中期反馈后，项目组围绕"段落
p_midterm = find_para_el(body, '中期反馈后，项目组围绕')
if p_midterm is not None:
    replace_run_text_in_para(
        p_midterm,
        '中期反馈后，项目组围绕测试大纲、自测报告、开发手册、UI设计文档、设备类型字典和支撑过程材料对系统需求、接口、数据和证据链进行了统一梳理。当前实现已形成"灾害数据导入-设备管理-模型训练-策略测试-场景回放-链路仿真-证据导出"的主流程，后端通过训练任务、策略测试任务、回放会话、专用设备数据库和灾害导入记录支撑前端操作留痕。',
        '2026年5月30日中期检查中，专家组就《广播通信融合网络模拟与资源优化调度仿真系统（原型系统）》提出了以下系统性整改意见：'
        '一是灾害场景边界定义不明确，需说明灾害类型、三断程度、损毁比例及恢复目标；'
        '二是用户节点规模口径不统一，1200、3200等数字在材料中前后表述不一致；'
        '三是多制式通信逻辑不清晰，700MHz、WiFi、卫星、短波各自在系统中的作用需明确；'
        '四是网络拓扑、设备组成、网关部署及AI训练推理位置描述不足；'
        '五是数据集来源主要集中于近期采集，历史统计类数据未充分利用，需区分受灾群众公网恢复需求与救援人员专网保障需求；'
        '六是算法对比验证不充分，需补充精度、计算量、训练时间等指标。'
        '中期反馈后，项目组围绕测试大纲、自测报告、开发手册、UI设计文档、设备类型字典和支撑过程材料对系统需求、接口、数据和证据链进行了统一梳理。'
        '当前实现已形成"灾害数据导入—设备管理—模型训练—策略测试—场景回放—链路仿真—证据导出"的主流程，'
        '后端通过训练任务、策略测试任务、回放会话、专用设备数据库和灾害导入记录支撑前端操作留痕。'
    )
    print('已更新背景段落')

# ── STEP 3: 在 2.2 分析目标末尾补充第（6）（7）条 ─────────────────────────────
p_goal5 = find_para_el(body, '（5）识别仍需补充确认的风险项')
if p_goal5 is not None:
    new6 = make_body_para(
        '（6）根据专家评审意见，对灾害场景边界（灾害类型、破坏对象、三断程度、损毁比例、空间范围）、'
        '节点规模口径（1200中期阈值/1500结项阈值/实际测试规模）、'
        '数据集来源与时间一致性、多制式通信逻辑、'
        '网络拓扑与AI部署位置及算法对比验证等内容开展影响范围分析，并纳入回归测试计划。'
    )
    new7 = make_body_para(
        '（7）识别受灾群众恢复公网通信需求与救援人员现场救援专网通信需求之间的差异，'
        '对系统功能布局、资源配置逻辑和测试用例覆盖范围的影响，'
        '确保需求分类在设计、测试和证据材料中一致体现。'
    )
    insert_after(body, p_goal5, new7)
    insert_after(body, p_goal5, new6)
    print('已补充分析目标（6）（7）')

# ── STEP 4: 在 3.3 之后新增 3.4 专家评审整改需求说明 ─────────────────────────
# 找到 H1 "4 影响范围分析"
h1_4 = find_para_el(body, '4 影响范围分析')
if h1_4 is not None and get_pStyle(h1_4) == '2':
    blocks = [
        make_h2('3.4 专家评审整改需求说明'),
        make_body_para(
            '2026年5月30日中期检查中，专家组就《广播通信融合网络模拟与资源优化调度仿真系统（原型系统）》'
            '提出了多项具体整改要求。本节将上述要求归类为需求变更并分析其影响范围，'
            '具体整改项目见表3-5。'
        ),
        make_body_para('表3-5 专家评审整改需求一览表'),
    ]
    for blk in reversed(blocks):
        insert_before(body, h1_4, blk)

    # 插入表3-5
    table_ref = find_para_el(body, '表3-5 专家评审整改需求一览表')
    if table_ref is not None:
        tbl = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(f'{{{NS}}}w', '9340'); tblW.set(f'{{{NS}}}type', 'dxa')
        tblPr.append(tblW)
        tblBorders = OxmlElement('w:tblBorders')
        for bn in ['top','left','bottom','right','insideH','insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(f'{{{NS}}}val','single'); b.set(f'{{{NS}}}sz','4')
            b.set(f'{{{NS}}}space','0'); b.set(f'{{{NS}}}color','000000')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        def cell(text, w, bold=False):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(f'{{{NS}}}w', str(w)); tcW.set(f'{{{NS}}}type', 'dxa')
            tcPr.append(tcW); tc.append(tcPr)
            p = make_body_para(text, indent=False)
            # 修正 bold
            for r in p.findall(f'{{{NS}}}r'):
                rpr = r.find(f'{{{NS}}}rPr')
                if rpr is not None:
                    b = rpr.find(f'{{{NS}}}b')
                    if b is not None:
                        if bold:
                            b.attrib.pop(f'{{{NS}}}val', None)
                        else:
                            b.set(f'{{{NS}}}val', '0')
            tc.append(p)
            return tc

        def row(cells):
            tr = OxmlElement('w:tr')
            for c in cells: tr.append(c)
            return tr

        col_w = [700, 2700, 1700, 2000, 2240]
        tbl.append(row([cell(h, w, bold=True) for h, w in zip(
            ['编号','整改内容','来源专家','影响范围','落实状态'], col_w)]))

        data = [
            ('ZJ-01',
             '明确灾害场景边界定义：灾害类型（暴雨/台风/地震/洪水）、三断程度、损毁比例、恢复目标和空间范围应在设计方案与测试大纲中统一规定。',
             '宋令阳/林树青', '总体设计方案、测试大纲、场景导入接口', '待补充'),
            ('ZJ-02',
             '统一用户节点规模口径：中期判定阈值≥1200，结项阈值≥1500，实际测试规模（3200/3500等）单独说明，避免在报告、日志和截图中混用。',
             '胡燕祝', '测试大纲、自测报告、回放日志、截图证据', '已部分落实'),
            ('ZJ-03',
             '区分受灾群众公网恢复需求与救援人员专网保障需求：两类需求的终端能力、通信方式和资源编排方式不同，应在需求模型、场景配置和算法设计中分类体现。',
             '冯伟', '多模态仿真环境、场景配置模块', '待设计'),
            ('ZJ-04',
             '补充网络拓扑、设备组成、网关部署及AI训练推理位置说明：需在总体设计和接口文档中说明不同计算节点的职责分工。',
             '杨冬', '总体设计方案、接口文档', '待补充'),
            ('ZJ-05',
             '补充多制式通信逻辑说明：700MHz、WiFi、卫星、短波各制式在系统中的作用及融合网关选路逻辑应在设备类型字典和测试用例中明确体现。',
             '黄宇红', '设备类型字典、测试大纲', '部分落实'),
            ('ZJ-06',
             '补充算法对比验证：在现有PPO/MAPPO等算法基础上，补充算法精度、计算量、训练时间、推理时间，及与传统方法或其他强化学习算法的对比结果。',
             '杨冬/黄宇红', '算法评估模块、测试大纲', '待补充'),
            ('ZJ-07',
             '补充数据集来源及时间一致性：融合历史统计类数据，统一数据采集时间范围表述，明确数据集本体与说明文档的关系。',
             '冯伟/尚红', '数据集、总体设计方案', '待补充'),
        ]
        for num, content, expert, scope, status in data:
            tbl.append(row([cell(v, w) for v, w in zip(
                [num, content, expert, scope, status], col_w)]))

        insert_after(body, table_ref, tbl)
        print('已插入 3.4 专家评审整改需求说明及表3-5')

# ── STEP 5: 新增 4.6 网络拓扑、AI部署与多制式通信影响分析 ─────────────────────
h1_5 = find_para_el(body, '5 回归测试范围分析')
if h1_5 is not None and get_pStyle(h1_5) == '2':
    blocks = [
        make_h2('4.6 网络拓扑、AI部署与多制式通信影响分析'),
        make_body_para(
            '根据专家评审意见，《广播通信融合网络模拟与资源优化调度仿真系统（原型系统）》'
            '在以下三个方面存在新增需求：'
            '一是网络拓扑结构、设备组成、网关部署位置和路由逻辑需要在设计方案和仿真系统中明确体现，'
            '不能仅强调AI算法而忽略组网属性；'
            '二是AI训练和推理的部署位置（中心平台、边缘设备、融合网关或其他计算节点）'
            '需要在总体架构和接口说明中明确；'
            '三是多制式通信逻辑（700MHz、WiFi、卫星、短波及融合网关选路）'
            '需要在设备类型字典和测试用例中清晰展现，'
            '并区分单独制式测试与多制式混合测试之间的关系。'
        ),
        make_body_para(
            '上述新增需求对原型系统的影响主要体现在以下四个方面：'
            '（1）场景数据载入时需支持网络拓扑和设备层次关系的输入，'
            '候选站点数据应能够表达不同制式设备的覆盖范围、部署约束和连接逻辑；'
            '（2）策略输出需包含AI推理位置和计算路径说明，'
            '便于测试人员核查策略生成过程与部署方案的一致性；'
            '（3）设备类型字典需进一步明确不同制式在融合网关中的路由和选择逻辑，'
            '以及700MHz广播通信方式在用户侧的覆盖和接收说明；'
            '（4）测试大纲中的第三方检测用例需覆盖多制式切换验证场景，'
            '并形成单制式与混合制式的独立证据材料。'
        ),
        make_body_para('表4-5 网络拓扑、AI部署与多制式通信影响分析表'),
    ]
    for blk in reversed(blocks):
        insert_before(body, h1_5, blk)
    print('已插入 4.6 节')

# ── STEP 6: 在 7.1 风险表中追加新风险行 ──────────────────────────────────────
# 找风险表 (第一个在"表7-1"之后的 tbl)
table71_ref = find_para_el(body, '表7-1 风险评估及应对措施表')
if table71_ref is not None:
    # 找紧接其后的 tbl
    children = list(body)
    risk_tbl = None
    for i, c in enumerate(children):
        if c is table71_ref:
            for j in range(i+1, len(children)):
                if children[j].tag == f'{{{NS}}}tbl':
                    risk_tbl = children[j]
                    break
            break

    if risk_tbl is not None:
        def risk_cell(text, w):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(f'{{{NS}}}w', str(w)); tcW.set(f'{{{NS}}}type', 'dxa')
            tcPr.append(tcW); tc.append(tcPr)
            tc.append(make_body_para(text, indent=False))
            return tc

        def risk_row(cells):
            tr = OxmlElement('w:tr')
            for c in cells: tr.append(c)
            return tr

        # 从现有表格的列宽推断
        new_risks = [
            ('R-08',
             '灾害场景边界表述不明确，两个灾害场景表现趋同，设置意义被削弱。',
             '专家评审时质疑场景代表性，无法支撑多灾害覆盖指标。',
             '在设计方案和测试大纲中给出灾害类型、三断程度、损毁比例和空间范围的明确定义；'
             '两个场景的受损特征、通信恢复方式和资源配置策略应体现差异。'),
            ('R-09',
             '多制式通信逻辑描述不清晰，700MHz、WiFi、卫星、短波责任边界模糊。',
             '单一制式测试与混合制式测试关系不明，第三方检测时证据不充分。',
             '在设备类型字典和测试用例中明确各制式在系统中的具体作用，'
             '分别设置单制式和混合制式测试场景，形成独立证据材料。'),
            ('R-10',
             '数据集时间表述前后不一致，PPT与文档对采集时间的描述存在差异。',
             '评审中对数据集实时性和覆盖范围产生争议，影响数据集指标判定。',
             '统一数据采集时间范围表述，补充历史统计类数据，'
             '明确数据集本体与说明文档的关系，对齐所有材料中的时间字段。'),
        ]
        # 从现有行推断列宽
        existing_rows = risk_tbl.findall(f'{{{NS}}}tr')
        col_widths = [800, 2500, 2500, 3540]  # 默认，与表7-1格式一致
        if existing_rows:
            first_row = existing_rows[0]
            tcs = first_row.findall(f'{{{NS}}}tc')
            ws = []
            for tc in tcs:
                tcW = tc.find(f'.//{{{NS}}}tcW')
                if tcW is not None:
                    try: ws.append(int(tcW.get(f'{{{NS}}}w', '0')))
                    except: ws.append(0)
            if len(ws) == 4:
                col_widths = ws

        for r_num, r_desc, r_impact, r_action in new_risks:
            risk_tbl.append(risk_row([
                risk_cell(v, w) for v, w in zip(
                    [r_num, r_desc, r_impact, r_action], col_widths)
            ]))
        print('已追加 R-08/R-09/R-10 风险行')

# ── STEP 7: 在结论段后追加整改总结段 ─────────────────────────────────────────
h1_9 = find_para_el(body, '9 附件及参考资料')
if h1_9 is not None and get_pStyle(h1_9) == '2':
    new_conclusion = make_body_para(
        '根据2026年5月30日专家评审意见，本次回归分析针对《广播通信融合网络模拟与资源优化调度仿真系统（原型系统）》'
        '新增了以下重点整改内容的影响范围评估：'
        '（1）灾害场景边界定义需在设计方案、测试大纲和证据材料中一致明确，'
        '两个灾害场景应体现受损特征、通信恢复方式和资源配置策略的差异；'
        '（2）用户节点规模口径已分为1200中期阈值、1500结项阈值和实际测试规模三层，'
        '需在所有相关文档和日志中保持一致；'
        '（3）多制式通信逻辑、网络拓扑、AI部署位置等新增需求已纳入第4.6节影响范围分析，'
        '并将在后续回归测试中设置专项验证用例；'
        '（4）受灾群众公网恢复需求与救援人员专网保障需求的分类差异已识别，'
        '需在仿真环境配置和测试用例设计中进一步体现。'
        '完成上述整改并通过高优先级回归测试验证后，'
        '系统具备进入中期自测复核、结项测试准备和第三方检测衔接的条件。'
    )
    insert_before(body, h1_9, new_conclusion)
    print('已追加结论补充段')

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(BASE + '13 需求回归分析_new.docx')
print(f'\n保存完成：{BASE}13 需求回归分析_new.docx')

# ── 验证 ──────────────────────────────────────────────────────────────────────
doc2 = docx.Document(BASE + '13 需求回归分析_new.docx')
print(f'段落数：{len(doc2.paragraphs)}，表格数：{len(doc2.tables)}')
print('\n所有标题：')
for i, p in enumerate(doc2.paragraphs):
    if p.style.name in ['Heading 1','Heading 2','Heading 3']:
        print(f'  [{i:3d}] {p.style.name}: {p.text}')
print('\n封面前5段：')
for i in range(5):
    print(f'  [{i}] {repr(doc2.paragraphs[i].text[:60])}')
