import docx
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def make_body_para_xml(text, indent_first=True):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(f'{{{NS}}}before', '0')
    spacing.set(f'{{{NS}}}after', '0')
    spacing.set(f'{{{NS}}}line', '360')
    spacing.set(f'{{{NS}}}lineRule', 'auto')
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(f'{{{NS}}}left', '0')
    ind.set(f'{{{NS}}}right', '0')
    if indent_first:
        ind.set(f'{{{NS}}}firstLine', '480')
    else:
        ind.set(f'{{{NS}}}firstLine', '0')
    pPr.append(ind)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(f'{{{NS}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS}}}eastAsia', '宋体')
    rFonts.set(f'{{{NS}}}cs', 'Times New Roman')
    rPr.append(rFonts)
    b = OxmlElement('w:b')
    b.set(f'{{{NS}}}val', '0')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(f'{{{NS}}}val', '24')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_heading1_xml(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(f'{{{NS}}}val', '2')
    pPr.append(pStyle)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(f'{{{NS}}}val', '0')
    numId = OxmlElement('w:numId')
    numId.set(f'{{{NS}}}val', '0')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    spacing = OxmlElement('w:spacing')
    spacing.set(f'{{{NS}}}before', '0')
    spacing.set(f'{{{NS}}}after', '0')
    spacing.set(f'{{{NS}}}line', '360')
    spacing.set(f'{{{NS}}}lineRule', 'auto')
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(f'{{{NS}}}left', '0')
    ind.set(f'{{{NS}}}right', '0')
    ind.set(f'{{{NS}}}firstLine', '0')
    pPr.append(ind)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(f'{{{NS}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS}}}eastAsia', '宋体')
    rFonts.set(f'{{{NS}}}cs', 'Times New Roman')
    rPr.append(rFonts)
    b_el = OxmlElement('w:b')
    rPr.append(b_el)
    sz = OxmlElement('w:sz')
    sz.set(f'{{{NS}}}val', '24')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_heading2_xml(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(f'{{{NS}}}val', '3')
    pPr.append(pStyle)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(f'{{{NS}}}val', '1')
    numId = OxmlElement('w:numId')
    numId.set(f'{{{NS}}}val', '0')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    spacing = OxmlElement('w:spacing')
    spacing.set(f'{{{NS}}}before', '60')
    spacing.set(f'{{{NS}}}after', '60')
    spacing.set(f'{{{NS}}}line', '360')
    spacing.set(f'{{{NS}}}lineRule', 'auto')
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(f'{{{NS}}}left', '0')
    ind.set(f'{{{NS}}}right', '0')
    ind.set(f'{{{NS}}}firstLine', '0')
    pPr.append(ind)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(f'{{{NS}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS}}}eastAsia', '宋体')
    rFonts.set(f'{{{NS}}}cs', 'Times New Roman')
    rPr.append(rFonts)
    b_el = OxmlElement('w:b')
    rPr.append(b_el)
    sz = OxmlElement('w:sz')
    sz.set(f'{{{NS}}}val', '24')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_after(body, ref_el, new_el):
    children = list(body)
    for i, child in enumerate(children):
        if child is ref_el:
            body.insert(i + 1, new_el)
            return i + 1
    return -1


def insert_before(body, ref_el, new_el):
    children = list(body)
    for i, child in enumerate(children):
        if child is ref_el:
            body.insert(i, new_el)
            return i
    return -1


def replace_para_text(para_el, new_text):
    """Replace the text in an existing paragraph element preserving its formatting"""
    # Remove all existing runs
    for r in para_el.findall(f'{{{NS}}}r'):
        para_el.remove(r)
    for bm in para_el.findall(f'{{{NS}}}bookmarkStart'):
        para_el.remove(bm)
    for bm in para_el.findall(f'{{{NS}}}bookmarkEnd'):
        para_el.remove(bm)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(f'{{{NS}}}ascii', 'Times New Roman')
    rFonts.set(f'{{{NS}}}hAnsi', 'Times New Roman')
    rFonts.set(f'{{{NS}}}eastAsia', '宋体')
    rFonts.set(f'{{{NS}}}cs', 'Times New Roman')
    rPr.append(rFonts)
    b = OxmlElement('w:b')
    b.set(f'{{{NS}}}val', '0')
    rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(f'{{{NS}}}val', '24')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = new_text
    r.append(t)
    para_el.append(r)


# ======== MAIN SCRIPT ========

doc = docx.Document('word/中期反馈/修改文件/13 需求回归分析_new.docx')
body = doc.element.body

paras = doc.paragraphs

# ---- 1. Update Background section (para 67) ----
old_p67 = paras[67]._p
new_bg = (
    "2026年5月30日，课题组在中期检查中收到专家组系统性意见，"
    "涉及需求文件整改、节点规模口径统一、灾害场景边界明确、"
    "数据集来源补充、多制式通信逻辑说明、网络拓扑与AI部署位置补充、"
    "算法评估对比补充、受灾群众与救援人员需求分类、"
    "以及测试大纲合并等多个方面。中期反馈后，项目组围绕测试大纲、"
    "自测报告、开发手册、UI设计文档、设备类型字典和支撑过程材料"
    "对系统需求、接口、数据和证据链进行了统一梳理。"
    "当前实现已形成“灾害数据导入-设备管理-模型训练-策略测试-"
    "场景回放-链路仳真-证据导出”的主流程，后端通过训练任务、"
    "策略测试任务、回放会话、专用设备数据库和灾害导入记录支撑前端操作留痕。"
)
replace_para_text(old_p67, new_bg)
print("Updated paragraph 67 (background section)")

# ---- 2. Update para 68 (third background paragraph) to add expert feedback mention ----
old_p68 = paras[68]._p
new_p68_text = (
    "为确认上述需求调整对系统功能、接口数据、测试用例、测试环境"
    "和验收证据的影响，同时将专家评审所提出的灾害场景边界、"
    "多制式通信、网络拓扑、AI部署位置、数据集、算法对比和需求分类"
    "等新增要求纳入需求基线并开展影响范围分析，"
    "特开展本次需求回归分析。本文档用于识别需求基线与当前实现之间"
    "的差异，明确需要回归验证的范围和优先级，建立需求、设计、"
    "开发、测试和验收证据之间的追踪关系，并为后续中期自测复核、"
    "结项测试准备和第三方检测衔接提供依据。"
)
replace_para_text(old_p68, new_p68_text)
print("Updated paragraph 68")

doc.save('word/中期反馈/修改文件/13 需求回归分析_new.docx')
print("Intermediate save OK")

# Reload to get fresh paragraph indices
doc = docx.Document('word/中期反馈/修改文件/13 需求回归分析_new.docx')
body = doc.element.body
paras = doc.paragraphs

# ---- 3. Add new goals to section 2.2 (after the last numbered goal) ----
# Find para with text starting with "（5）识别仍需补充"
target_5 = None
for i, p in enumerate(paras):
    if p.text.startswith('（5）识别仍需补充'):
        target_5 = p._p
        print(f"Found goal 5 at index {i}: {p.text[:50]}")
        break

if target_5 is not None:
    new_goals = [
        "（6）根据专家评审意见，对灾害场景边界（灾害类型、破坏对象、三断程度、损毁比例和空间范围）、"
        "节点规模口径统一（1200/1500/3200/3500区分）、数据集来源与时间一致性、"
        "多制式通信逻辑说明、网络拓扑与AI布署位置补充、算法对比验证"
        "等新增内容开展影响范围分析，并将其纳入回归测试计划。",

        "（7）识别受灾群众恢复公网通信需求与救援人员现场救援专网通信需求"
        "之间的差异对系统功能布局和资源配置方式的影响，确保需求分类"
        "在设计、测试和证据材料中一致体现。",
    ]
    current_ref = target_5
    for txt in new_goals:
        new_p = make_body_para_xml(txt)
        idx = insert_after(body, current_ref, new_p)
        current_ref = new_p

doc.save('word/中期反馈/修改文件/13 需求回归分析_new.docx')
print("Saved after adding goals 6 and 7")

# Reload
doc = docx.Document('word/中期反馈/修改文件/13 需求回归分析_new.docx')
body = doc.element.body
paras = doc.paragraphs

# ---- 4. Add new subsection 3.4 "专家评审整改需求说明" after 3.3 ----
# Find the Heading 1 "4 影响范围分析"
h1_4_idx = None
for i, p in enumerate(paras):
    if p.style.name == 'Heading 1' and p.text.startswith('4 影响范围'):
        h1_4_idx = i
        print(f"Found H1 '4 影响范围分析' at index {i}")
        break

if h1_4_idx is not None:
    h1_4_el = paras[h1_4_idx]._p

    new_subsec_content = [
        # Heading 2
        ('h2', '3.4 专家评审整改需求说明'),
        # paragraphs
        ('body', '2026年5月30日中期检查中，专家组提出了多项具体整改要求。'
                 '本节将上述要求将其归类为需求变更并进行影响范围分析，'
                 '具体整改项目如下表3-5所示。'),
        ('table', '表3-5 专家评审整改需求一览表'),
    ]

    current_ref = paras[h1_4_idx - 1]._p  # para before H1 "4"
    # We'll insert before h1_4_el
    for item_type, item_text in reversed(new_subsec_content):
        if item_type == 'h2':
            new_el = make_heading2_xml(item_text)
        elif item_type == 'body':
            new_el = make_body_para_xml(item_text)
        else:
            new_el = make_body_para_xml(item_text)
        insert_before(body, h1_4_el, new_el)

    # Now add the actual table with expert feedback items
    # We'll add a proper table after the "表3-5" reference para
    # First find the newly inserted "表3-5" para
    table_ref_el = None
    for child in list(body):
        if child.tag == f'{{{NS}}}p':
            t_els = child.findall(f'.//{{{NS}}}t')
            text = ''.join(t.text or '' for t in t_els)
            if '表3-5' in text:
                table_ref_el = child
                break

    if table_ref_el is not None:
        # Build a simple table for expert feedback requirements
        tbl = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(f'{{{NS}}}w', '9340')
        tblW.set(f'{{{NS}}}type', 'dxa')
        tblPr.append(tblW)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(f'{{{NS}}}val', 'single')
            border.set(f'{{{NS}}}sz', '4')
            border.set(f'{{{NS}}}space', '0')
            border.set(f'{{{NS}}}color', '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        def make_cell(text, bold=False, width=None):
            tc = OxmlElement('w:tc')
            if width:
                tcPr = OxmlElement('w:tcPr')
                tcW = OxmlElement('w:tcW')
                tcW.set(f'{{{NS}}}w', str(width))
                tcW.set(f'{{{NS}}}type', 'dxa')
                tcPr.append(tcW)
                tc.append(tcPr)
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            spacing = OxmlElement('w:spacing')
            spacing.set(f'{{{NS}}}before', '0')
            spacing.set(f'{{{NS}}}after', '0')
            spacing.set(f'{{{NS}}}line', '360')
            spacing.set(f'{{{NS}}}lineRule', 'auto')
            pPr.append(spacing)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{NS}}}left', '0')
            ind.set(f'{{{NS}}}right', '0')
            ind.set(f'{{{NS}}}firstLine', '0')
            pPr.append(ind)
            p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(f'{{{NS}}}ascii', 'Times New Roman')
            rFonts.set(f'{{{NS}}}hAnsi', 'Times New Roman')
            rFonts.set(f'{{{NS}}}eastAsia', '宋体')
            rFonts.set(f'{{{NS}}}cs', 'Times New Roman')
            rPr.append(rFonts)
            if bold:
                b_el = OxmlElement('w:b')
                rPr.append(b_el)
            else:
                b_el = OxmlElement('w:b')
                b_el.set(f'{{{NS}}}val', '0')
                rPr.append(b_el)
            sz = OxmlElement('w:sz')
            sz.set(f'{{{NS}}}val', '24')
            rPr.append(sz)
            r.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = text
            r.append(t_el)
            p.append(r)
            tc.append(p)
            return tc

        def make_row(cells_data):
            tr = OxmlElement('w:tr')
            for text, bold, width in cells_data:
                tr.append(make_cell(text, bold=bold, width=width))
            return tr

        # Header row
        header = make_row([
            ('编号', True, 800),
            ('整改项目', True, 3200),
            ('来源专家', True, 1800),
            ('影响范围', True, 1800),
            ('落实状态', True, 1740),
        ])
        tbl.append(header)

        rows_data = [
            ('ZJ-01', '将组网方案报告明确标注为“初稿”或“初版完成”，避免表述为最终完成稿。',
             '胡燕祝教授', '天站级文档、自测报告', '待更新'),
            ('ZJ-02', '统一1200/1500/3200/3500节点口径：中期判定阈値≥ 1200，结项阈値≥ 1500，实际测试规模单独说明。',
             '胡燕祝教授', '测试大纲、自测报告、回放日志', '已部分落实'),
            ('ZJ-03', '明确灾害场景边界：灾害类型、破坏对象、三断程度、损毁比例及恢复目标应在设计方案和测试大纲中给出明确定义。',
             '宋令阳教授', '场景设计、测试大纲', '待补充'),
            ('ZJ-04', '将灾害通信需求分为两类：受灾群众公网恢复需求与救援人员专网保障需求，在模型、算法和仳真系统中分别设置。',
             '冯伟教授', '多模态仳真环境、算法模块', '待设计'),
            ('ZJ-05', '补充网络拓扑、设备组成、网关部署位置及AI训练/推理部署位置说明。',
             '杨冬教授', '组网方案、系统设计文档', '待补充'),
            ('ZJ-06', '补充多制式通信逻辑：700MHz、WiFi、卫星、短波各制式及融合网关作用应在设备类型字典和测试用例中明确说明。',
             '黄宇红院长', '设备类型字典、测试大纲', '部分落实'),
            ('ZJ-07', '补充算法对比验证：补充算法精度、计算量、训练时间、推理时间及与其他算法对比结果。',
             '杨冬教授/黄宇红', '算法评估模块、测试大纲', '待补充'),
            ('ZJ-08', '补充数据集来源、历史数据和采集时间一致性说明，隐设幾个字段和数据集本体及说明文档关系。',
             '冯伟教授/尚红主任', '数据集、设计文档', '待补充'),
        ]

        for row_data in rows_data:
            num, change, expert, scope, status = row_data
            tr = make_row([
                (num, False, 800),
                (change, False, 3200),
                (expert, False, 1800),
                (scope, False, 1800),
                (status, False, 1740),
            ])
            tbl.append(tr)

        # Insert table after the "表3-5" para
        insert_after(body, table_ref_el, tbl)
        print("Inserted expert feedback requirements table")

doc.save('word/中期反馈/修改文件/13 需求回归分析_new.docx')
print("Saved after section 3.4")

# ---- 5. Add new subsection 4.6 for network topology and AI deployment impact ----
doc = docx.Document('word/中期反馈/修改文件/13 需求回归分析_new.docx')
body = doc.element.body
paras = doc.paragraphs

# Find H1 "5 回归测试范围分析"
h1_5_el = None
for i, p in enumerate(paras):
    if p.style.name == 'Heading 1' and '5 回归测试范围' in p.text:
        h1_5_el = p._p
        print(f"Found H1 '5 回归测试范围分析' at index {i}")
        break

if h1_5_el is not None:
    new_sec46 = [
        ('h2', '4.6 网络拓扑、AI部署与多制式通信影响分析'),
        ('body',
         '根据专家评审意见，系统在以下三个方面存在新增需求：'
         '一是网络拓扑结构、设备组成、网关部署位置和路由逻辑'
         '需要在设计方案和仳真系统中明确体现；'
         '二是AI训练和推理的部署位置（中心平台、边缘设备、融合网关或其他计算节点）'
         '需要在整体设计和接口说明中体现；'
         '三是多制式通信逻辑（700MHz、WiFi、卫星、短波及融合网关选路）'
         '需要在设备类型字典和测试用例中清晰展现。'
        ),
        ('body',
         '上述新增需求对原型系统的影响主要体现在：（1）场景数据载入'
         '时需支持网络拓扑和设备层次关系输入；（2）策略输出需包含'
         'AI推理位置和计算路径说明；（3）设备类型字典需进一步明确不同'
         '制式在融合网关中的路由和选择逻辑；（4）测试大纲中的第三方'
         '检测用例需覆盖多制式切换验证场景。'
        ),
        ('body', '表4-5 网络拓扑、AI部署与多制式影响分析表'),
    ]

    for item_type, item_text in reversed(new_sec46):
        if item_type == 'h2':
            new_el = make_heading2_xml(item_text)
        else:
            new_el = make_body_para_xml(item_text)
        insert_before(body, h1_5_el, new_el)

    print("Inserted section 4.6")

doc.save('word/中期反馈/修改文件/13 需求回归分析_new.docx')
print("Saved after 4.6")

# ---- 6. Add new risk items to section 7.1 ----
doc = docx.Document('word/中期反馈/修改文件/13 需求回归分析_new.docx')
body = doc.element.body
paras = doc.paragraphs

# Find "表7-1" and get the table right after it
table71_ref = None
table71_idx = None
for i, p in enumerate(paras):
    if '表7-1' in p.text:
        table71_ref = p._p
        table71_idx = i
        print(f"Found 表7-1 ref at index {i}")
        break

if table71_ref is not None:
    # The risk table is in doc.tables - find which one is after 表7-1
    # We need to find the table element after table71_ref in body
    found_table = False
    for child in list(body):
        if child is table71_ref:
            found_table = True
            continue
        if found_table and child.tag == f'{{{NS}}}tbl':
            risk_table = child
            print("Found risk table after 表7-1")
            # Add 3 new rows to this table
            new_risks = [
                ('R-08', '灾害场景边界表述不明确，导致不同灾害场景寻找相同论证话。',
                 '可能导致两个场景设置意义被削弱，专家评审时质疑。',
                 '在设计方案和测试大纲中给出灾害类型、三断程度、损毁比例和空间范围的明确定义。'),
                ('R-09', '多制式通信逻辑表述不清晰，导致700MHz、WiFi、卫星、短波责任划分模糊。',
                 '将出现单一制式疗照多制式说明的风险，单经品测试与混合测试关系不明。',
                 '在设备类型字典和测试用例中明确说明各制式在系统中的具体作用，并设计单独制式和混合制式测试用例。'),
                ('R-10', '数据集时间表述不一致，导致考核时对数据实时性产生质疑。',
                 '将导致中期评审中对数据集质量和时间新鲜性的争议。',
                 '全面梳理数据采集时间范围，合并历史统计数据和现场采集数据，对齐PPT和文档中的时间表述。'),
            ]
            for r_data in new_risks:
                r_num, r_desc, r_impact, r_action = r_data
                tr = OxmlElement('w:tr')
                for text, width in [(r_num, 800), (r_desc, 2500), (r_impact, 2500), (r_action, 3540)]:
                    tc = OxmlElement('w:tc')
                    tcPr = OxmlElement('w:tcPr')
                    tcW = OxmlElement('w:tcW')
                    tcW.set(f'{{{NS}}}w', str(width))
                    tcW.set(f'{{{NS}}}type', 'dxa')
                    tcPr.append(tcW)
                    tc.append(tcPr)
                    p_el = make_body_para_xml(text, indent_first=False)
                    tc.append(p_el)
                    tr.append(tc)
                risk_table.append(tr)
            print("Added 3 new risk rows")
            break

doc.save('word/中期反馈/修改文件/13 需求回归分析_new.docx')
print("Saved after risk table update")

# ---- 7. Update conclusion section ----
doc = docx.Document('word/中期反馈/修改文件/13 需求回归分析_new.docx')
body = doc.element.body
paras = doc.paragraphs

# Find last paragraph of conclusion (before H1 "9 附件")
h1_9_el = None
for i, p in enumerate(paras):
    if p.style.name == 'Heading 1' and '9 附件' in p.text:
        h1_9_el = p._p
        print(f"Found H1 '9 附件' at index {i}")
        break

if h1_9_el is not None:
    new_conclusion_para = make_body_para_xml(
        "根据2026年5月30日专家评审意见，"
        "本次回归分析新增了以下重点整改内容的影响范围评估："
        "（1）灾害场景边界定义需要在设计方案、测试大纲和证据材料中一致明确，"
        "涵盖灾害类型、三断程度、损毁比例和恢复目标；"
        "（2）用户节点规模口径已分为1200中期阈値、1500结项阈値和实际测试规模三层说明，"
        "需在所有相关文档和日志中保持一致；"
        "（3）多制式通信逻辑、网络拓扑、AI部署位置和数据集来源等新增需求"
        "已纳入影响范围分析和回归测试计划；"
        "（4）受灾群众公网需求与救援人员专网需求的分类影响已识别，"
        "需在后续测试设计和仳真环境配置中进一步体现。"
        "在完成上述整改并通过高优先级回归测试验证后，"
        "系统具备进入中期自测复核、结项测试准备和第三方检测衔接的条件。"
    )
    insert_before(body, h1_9_el, new_conclusion_para)
    print("Added conclusion paragraph")

doc.save('word/中期反馈/修改文件/13 需求回归分析_new.docx')
print("All modifications saved!")
print("Output: word/中期反馈/修改文件/13 需求回归分析_new.docx")
